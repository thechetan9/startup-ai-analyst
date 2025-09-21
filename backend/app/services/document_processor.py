"""
Document processing service for PDF and DOCX files
Handles extraction, parsing, and structuring of startup documents
"""

import os
import io
import logging
from typing import Dict, Any, Optional, List, Tuple
from pathlib import Path
import asyncio
import aiofiles
from datetime import datetime

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
from google.cloud import vision, storage
import google.generativeai as genai

from app.core.config import settings
from app.models.documents import (
    DocumentType, ProcessingStatus, ExtractedContent, 
    PitchDeckData, FinancialStatementData, ProcessingResult
)

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document processing for various file types"""
    
    def __init__(self):
        self.vision_client = vision.ImageAnnotatorClient()
        self.storage_client = storage.Client()
        genai.configure(api_key=settings.gemini_api_key)
        self.gemini_model = genai.GenerativeModel('gemini-pro')
    
    async def process_document(
        self,
        file_path: str,
        document_type: DocumentType,
        startup_id: str,
        progress_callback=None
    ) -> ProcessingResult:
        """Main document processing entry point with progress tracking"""

        start_time = asyncio.get_event_loop().time()
        document_id = f"{startup_id}_{Path(file_path).stem}"

        try:
            # Check file size first
            file_size = Path(file_path).stat().st_size
            max_size = 50 * 1024 * 1024  # 50MB limit

            if file_size > max_size:
                error_msg = f"File too large: {file_size / 1024 / 1024:.1f}MB. Maximum allowed: {max_size / 1024 / 1024}MB"
                if progress_callback:
                    await progress_callback(error_msg, -1, "error")
                raise ValueError(error_msg)

            if progress_callback:
                await progress_callback("Validating file format...", 5, "processing")

            # Determine file type and process accordingly
            file_extension = Path(file_path).suffix.lower()

            if progress_callback:
                await progress_callback("Starting content extraction...", 10, "processing")

            if file_extension == '.pdf':
                extracted_content = await self._process_pdf(file_path, document_type, progress_callback)
            elif file_extension in ['.docx', '.doc']:
                extracted_content = await self._process_docx(file_path, document_type, progress_callback)
            else:
                error_msg = f"Unsupported file type: {file_extension}"
                if progress_callback:
                    await progress_callback(error_msg, -1, "error")
                raise ValueError(error_msg)

            if progress_callback:
                await progress_callback("Structuring extracted data...", 70, "processing")

            # Structure the extracted data based on document type
            structured_data = await self._structure_data(
                extracted_content, document_type, startup_id, progress_callback
            )

            if progress_callback:
                await progress_callback("Finalizing results...", 95, "processing")

            processing_time = asyncio.get_event_loop().time() - start_time

            if progress_callback:
                await progress_callback("Processing completed successfully!", 100, "completed")

            return ProcessingResult(
                document_id=document_id,
                processing_status=ProcessingStatus.COMPLETED,
                extracted_content=extracted_content,
                structured_data=structured_data,
                processing_time_seconds=processing_time,
                processed_at=datetime.now()
            )

        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            processing_time = asyncio.get_event_loop().time() - start_time

            if progress_callback:
                await progress_callback(f"Processing failed: {str(e)}", -1, "error")

            return ProcessingResult(
                document_id=document_id,
                processing_status=ProcessingStatus.FAILED,
                processing_errors=[str(e)],
                processing_time_seconds=processing_time,
                processed_at=datetime.now()
            )
    
    async def _process_pdf(self, file_path: str, document_type: DocumentType, progress_callback=None) -> ExtractedContent:
        """Process PDF files using PyPDF2 and Google Cloud Vision"""

        try:
            if progress_callback:
                await progress_callback("Extracting text from PDF...", 20, "processing")

            # First, try text extraction with PyPDF2
            text_content = await self._extract_pdf_text(file_path)

            if progress_callback:
                await progress_callback("Analyzing text quality...", 35, "processing")

            # If text extraction is poor, use OCR with Cloud Vision
            if len(text_content.strip()) < 100:  # Threshold for poor text extraction
                logger.info("PDF has poor text extraction, using OCR")
                if progress_callback:
                    await progress_callback("Using OCR for better text extraction...", 45, "processing")
                text_content = await self._ocr_pdf_with_vision(file_path)

            if progress_callback:
                await progress_callback("Extracting document metadata...", 55, "processing")

            # Extract metadata
            metadata = await self._extract_pdf_metadata(file_path)

            if progress_callback:
                await progress_callback("PDF processing completed", 60, "processing")

            return ExtractedContent(
                document_id=Path(file_path).stem,
                raw_text=text_content,
                metadata=metadata,
                confidence_score=0.9 if len(text_content) > 100 else 0.6
            )

        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            if progress_callback:
                await progress_callback(f"PDF processing failed: {str(e)}", -1, "error")
            raise
    
    async def _process_docx(self, file_path: str, document_type: DocumentType, progress_callback=None) -> ExtractedContent:
        """Process DOCX files using python-docx"""
        
        try:
            if progress_callback:
                await progress_callback("Reading DOCX file...", 25, "processing")

            async with aiofiles.open(file_path, 'rb') as file:
                content = await file.read()

            if progress_callback:
                await progress_callback("Parsing document structure...", 30, "processing")

            # Parse DOCX content
            doc = DocxDocument(io.BytesIO(content))

            if progress_callback:
                await progress_callback("Extracting paragraphs...", 40, "processing")

            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())

            if progress_callback:
                await progress_callback("Extracting tables...", 50, "processing")

            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                tables_text.append("\n".join(table_data))
            
            # Combine all text
            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\nTABLES:\n" + "\n\n".join(tables_text)
            
            # Extract metadata
            metadata = {
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "word_count": len(full_text.split()),
                "has_images": len(doc.inline_shapes) > 0
            }
            
            return ExtractedContent(
                document_id=Path(file_path).stem,
                raw_text=full_text,
                metadata=metadata,
                confidence_score=0.95  # DOCX extraction is usually very reliable
            )
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            raise
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using PyPDF2"""
        
        text_content = []
        
        async with aiofiles.open(file_path, 'rb') as file:
            content = await file.read()
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        
        for page in pdf_reader.pages:
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(page_text)
            except Exception as e:
                logger.warning(f"Failed to extract text from page: {e}")
                continue
        
        return "\n\n".join(text_content)
    
    async def _ocr_pdf_with_vision(self, file_path: str) -> str:
        """Use Google Cloud Vision for OCR on PDF"""
        
        try:
            # Convert PDF to images and process with Vision API
            # This is a simplified version - in production, you'd want to
            # convert PDF pages to images first
            
            async with aiofiles.open(file_path, 'rb') as file:
                content = await file.read()
            
            # For now, we'll use the document text detection
            # In a full implementation, you'd convert PDF to images first
            image = vision.Image(content=content)
            response = self.vision_client.document_text_detection(image=image)
            
            if response.full_text_annotation:
                return response.full_text_annotation.text
            else:
                return ""
                
        except Exception as e:
            logger.error(f"OCR processing failed: {e}")
            return ""
    
    async def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF"""
        
        try:
            async with aiofiles.open(file_path, 'rb') as file:
                content = await file.read()
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            
            metadata = {
                "page_count": len(pdf_reader.pages),
                "file_size": len(content),
                "has_encryption": pdf_reader.is_encrypted
            }
            
            # Extract document info if available
            if pdf_reader.metadata:
                doc_info = pdf_reader.metadata
                metadata.update({
                    "title": doc_info.get("/Title", ""),
                    "author": doc_info.get("/Author", ""),
                    "subject": doc_info.get("/Subject", ""),
                    "creator": doc_info.get("/Creator", ""),
                    "producer": doc_info.get("/Producer", ""),
                    "creation_date": str(doc_info.get("/CreationDate", "")),
                    "modification_date": str(doc_info.get("/ModDate", ""))
                })
            
            return metadata
            
        except Exception as e:
            logger.error(f"Metadata extraction failed: {e}")
            return {"error": str(e)}
    
    async def _structure_data(
        self, 
        extracted_content: ExtractedContent, 
        document_type: DocumentType,
        startup_id: str
    ) -> Dict[str, Any]:
        """Structure extracted data based on document type using Gemini AI"""
        
        try:
            # Create document-type specific prompts
            prompt = self._get_structuring_prompt(document_type, extracted_content.raw_text)
            
            # Use Gemini to structure the data
            response = await self._query_gemini(prompt)
            
            # Parse and validate the response
            structured_data = self._parse_gemini_response(response, document_type)
            
            return structured_data
            
        except Exception as e:
            logger.error(f"Data structuring failed: {e}")
            return {"error": str(e), "raw_text_preview": extracted_content.raw_text[:500]}
    
    def _get_structuring_prompt(self, document_type: DocumentType, text: str) -> str:
        """Generate appropriate prompt for document structuring"""
        
        base_prompt = f"""
        Analyze the following {document_type.value} document and extract structured information.
        Return the response in JSON format with clear field names and values.
        
        Document content:
        {text[:4000]}  # Limit text to avoid token limits
        
        """
        
        if document_type == DocumentType.PITCH_DECK:
            return base_prompt + """
            Extract the following information:
            - company_name: Company name
            - tagline: Company tagline or mission
            - problem_statement: Problem being solved
            - solution_description: Solution description
            - market_size: Market size information (TAM, SAM, SOM)
            - business_model: Business model description
            - traction_metrics: Key traction metrics and numbers
            - financial_projections: Revenue projections and financial data
            - team_info: Founder and team information
            - funding_ask: Funding amount requested
            - use_of_funds: How funds will be used
            - competition_analysis: Competitive landscape
            """
        
        elif document_type == DocumentType.FINANCIAL_STATEMENT:
            return base_prompt + """
            Extract financial data:
            - period: Financial period (year, quarter)
            - revenue: Total revenue
            - gross_profit: Gross profit
            - operating_expenses: Operating expenses
            - net_income: Net income
            - cash_flow: Cash flow information
            - key_metrics: Important financial ratios and metrics
            """
        
        else:
            return base_prompt + """
            Extract key information relevant to startup evaluation:
            - key_points: Main points and insights
            - metrics: Any numerical data or KPIs
            - concerns: Potential red flags or concerns
            - opportunities: Growth opportunities mentioned
            """
    
    async def _query_gemini(self, prompt: str) -> str:
        """Query Gemini AI model"""
        
        try:
            response = self.gemini_model.generate_content(prompt)
            return response.text
        except Exception as e:
            logger.error(f"Gemini query failed: {e}")
            raise
    
    def _parse_gemini_response(self, response: str, document_type: DocumentType) -> Dict[str, Any]:
        """Parse and validate Gemini response"""
        
        try:
            # Try to extract JSON from the response
            import json
            import re
            
            # Look for JSON content in the response
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                return json.loads(json_str)
            else:
                # If no JSON found, create a structured response
                return {
                    "raw_analysis": response,
                    "document_type": document_type.value,
                    "parsing_note": "Could not extract structured JSON, returning raw analysis"
                }
                
        except Exception as e:
            logger.error(f"Response parsing failed: {e}")
            return {
                "error": str(e),
                "raw_response": response[:1000]
            }


# Global document processor instance
document_processor = DocumentProcessor()
