"""
Large file processing service with chunking and streaming
"""

import os
import io
import logging
from typing import Dict, Any, Optional, List, AsyncGenerator
from pathlib import Path
import asyncio
import aiofiles
from datetime import datetime

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
import tempfile

logger = logging.getLogger(__name__)

class LargeFileProcessor:
    """Handle large files with chunking and streaming"""
    
    def __init__(self, chunk_size: int = 10 * 1024 * 1024):  # 10MB chunks
        self.chunk_size = chunk_size
        self.max_file_size = 200 * 1024 * 1024  # 200MB absolute limit
    
    async def process_large_file(
        self, 
        file_path: str, 
        progress_callback=None
    ) -> Dict[str, Any]:
        """Process large files in chunks"""
        
        file_size = Path(file_path).stat().st_size
        
        if file_size > self.max_file_size:
            error_msg = f"File too large: {file_size / 1024 / 1024:.1f}MB. Maximum allowed: {self.max_file_size / 1024 / 1024}MB"
            if progress_callback:
                await progress_callback(error_msg, -1, "error")
            raise ValueError(error_msg)
        
        if progress_callback:
            await progress_callback("Preparing large file processing...", 5, "processing")
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return await self._process_large_pdf(file_path, progress_callback)
        elif file_extension in ['.docx', '.doc']:
            return await self._process_large_docx(file_path, progress_callback)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    async def _process_large_pdf(self, file_path: str, progress_callback=None) -> Dict[str, Any]:
        """Process large PDF files page by page"""
        
        try:
            if progress_callback:
                await progress_callback("Opening large PDF file...", 10, "processing")
            
            async with aiofiles.open(file_path, 'rb') as file:
                content = await file.read()
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            total_pages = len(pdf_reader.pages)
            
            if progress_callback:
                await progress_callback(f"Processing {total_pages} pages...", 15, "processing")
            
            extracted_text = []
            processed_pages = 0
            
            # Process pages in batches to avoid memory issues
            batch_size = 10
            for i in range(0, total_pages, batch_size):
                batch_end = min(i + batch_size, total_pages)
                
                # Process batch of pages
                batch_text = []
                for page_num in range(i, batch_end):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text.strip():
                            batch_text.append(f"--- Page {page_num + 1} ---\n{page_text}")
                        processed_pages += 1
                        
                        # Update progress
                        progress = 15 + (processed_pages / total_pages) * 50
                        if progress_callback:
                            await progress_callback(
                                f"Processed page {processed_pages}/{total_pages}", 
                                int(progress), 
                                "processing"
                            )
                        
                        # Small delay to prevent blocking
                        if page_num % 5 == 0:
                            await asyncio.sleep(0.01)
                            
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
                
                extracted_text.extend(batch_text)
                
                # Memory cleanup
                if i % 50 == 0:  # Every 50 pages
                    await asyncio.sleep(0.1)  # Allow other tasks to run
            
            if progress_callback:
                await progress_callback("Combining extracted text...", 70, "processing")
            
            full_text = "\n\n".join(extracted_text)
            
            # Extract metadata
            metadata = {
                "total_pages": total_pages,
                "processed_pages": processed_pages,
                "file_size_mb": Path(file_path).stat().st_size / 1024 / 1024,
                "extraction_method": "chunked_processing",
                "word_count": len(full_text.split()),
                "char_count": len(full_text)
            }
            
            if progress_callback:
                await progress_callback("Large PDF processing completed", 75, "processing")
            
            return {
                "text": full_text,
                "metadata": metadata,
                "confidence_score": 0.85,  # Slightly lower for large files
                "processing_method": "chunked"
            }
            
        except Exception as e:
            logger.error(f"Large PDF processing failed: {e}")
            if progress_callback:
                await progress_callback(f"Large PDF processing failed: {str(e)}", -1, "error")
            raise
    
    async def _process_large_docx(self, file_path: str, progress_callback=None) -> Dict[str, Any]:
        """Process large DOCX files section by section"""
        
        try:
            if progress_callback:
                await progress_callback("Opening large DOCX file...", 10, "processing")
            
            async with aiofiles.open(file_path, 'rb') as file:
                content = await file.read()
            
            if progress_callback:
                await progress_callback("Parsing document structure...", 20, "processing")
            
            doc = DocxDocument(io.BytesIO(content))
            
            total_paragraphs = len(doc.paragraphs)
            total_tables = len(doc.tables)
            
            if progress_callback:
                await progress_callback(f"Processing {total_paragraphs} paragraphs and {total_tables} tables...", 25, "processing")
            
            # Process paragraphs in batches
            paragraphs = []
            processed_paragraphs = 0
            batch_size = 50
            
            for i in range(0, total_paragraphs, batch_size):
                batch_end = min(i + batch_size, total_paragraphs)
                
                for para_idx in range(i, batch_end):
                    paragraph = doc.paragraphs[para_idx]
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text.strip())
                    processed_paragraphs += 1
                    
                    # Update progress
                    progress = 25 + (processed_paragraphs / total_paragraphs) * 30
                    if progress_callback and processed_paragraphs % 10 == 0:
                        await progress_callback(
                            f"Processed paragraph {processed_paragraphs}/{total_paragraphs}", 
                            int(progress), 
                            "processing"
                        )
                
                # Small delay to prevent blocking
                await asyncio.sleep(0.01)
            
            if progress_callback:
                await progress_callback("Processing tables...", 60, "processing")
            
            # Process tables
            tables_text = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                tables_text.append(f"--- Table {table_idx + 1} ---\n" + "\n".join(table_data))
                
                if progress_callback and table_idx % 5 == 0:
                    progress = 60 + (table_idx / total_tables) * 10
                    await progress_callback(
                        f"Processed table {table_idx + 1}/{total_tables}", 
                        int(progress), 
                        "processing"
                    )
            
            if progress_callback:
                await progress_callback("Combining extracted content...", 75, "processing")
            
            # Combine all text
            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\nTABLES:\n" + "\n\n".join(tables_text)
            
            # Extract metadata
            metadata = {
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "file_size_mb": Path(file_path).stat().st_size / 1024 / 1024,
                "extraction_method": "chunked_processing",
                "word_count": len(full_text.split()),
                "char_count": len(full_text),
                "has_images": len(doc.inline_shapes) > 0
            }
            
            if progress_callback:
                await progress_callback("Large DOCX processing completed", 80, "processing")
            
            return {
                "text": full_text,
                "metadata": metadata,
                "confidence_score": 0.9,  # DOCX is usually reliable
                "processing_method": "chunked"
            }
            
        except Exception as e:
            logger.error(f"Large DOCX processing failed: {e}")
            if progress_callback:
                await progress_callback(f"Large DOCX processing failed: {str(e)}", -1, "error")
            raise

# Global instance
large_file_processor = LargeFileProcessor()
